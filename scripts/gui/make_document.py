from docx           import Document
#from docx.oxml      import OxmlElement
#from docx.oxml.ns   import qn
#from docx2pdf       import convert
from docx.shared    import Cm


"""
TODO
- Replace all path to absolute path     : Pending
"""

def make_document(
        meta_data       = {},
        template_path   = "../../documents/template.docx",
        output_path     = "../../documents/checklist_report.docx",
        save_pdf        = False,
):
    
    document = Document(template_path)

    # set header information
    header          = document.sections[0].header

    header_data     = {
        "{{actuator_serial_no}}" : meta_data['actuator_serial_number'] + "           ",
        "{{doc_no}}"             : meta_data['document_number'],
        "{{rev no}}"             : "                   " + meta_data['revision_number']
    }

    for paragraph in header.paragraphs:
        for key,value in header_data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key,value)

    # set paragraph information
    main_paragraph_data     = {
        "{{date}}"                  : meta_data['date'],
        "{{assembled_by}}"          : meta_data['assembler_name'],
        "{{assembled_date}}"        : meta_data['assembly_date'],
        "{{assembler_signature}}"   : meta_data['assembler_signature'],
        "{{tested_by}}"             : meta_data['tester_name'],
        "{{tested_date}}"           : meta_data['testing_date'],
        "{{tester_signature}}"      : meta_data['tester_signature'],
        "{{approved_by}}"           : meta_data['approver_name'], 
        "{{approved_date}}"         : meta_data['approval_date'],
        "{{approver_signature}}"    : meta_data['approver_signature'],
        "{{end_remarks}}"           : meta_data['end_remarks']
    }

    main_paragraph_data = fix_spaces(main_paragraph_data)

    for i,paragraph in enumerate(document.paragraphs):
        for key,value in main_paragraph_data.items():
            if key in paragraph.text:
                paragraph.text      = paragraph.text.replace(key,value)

    # set table information
    check_box_status = {
            "{{ch_1}}" : meta_data['ch_1'],
            "{{ch_2}}" : meta_data['ch_2'],
            "{{ch_3}}" : meta_data['ch_3'],
            "{{ch_4}}" : meta_data['ch_4'],
            "{{ch_5}}" : meta_data['ch_5'],
            "{{ch_6}}" : meta_data['ch_6'],
            "{{ch_7}}" : meta_data['ch_7'],
            "{{ch_8}}" : meta_data['ch_8'],
        }

    yes_no_status   = {
        "{{yn_1}}" : meta_data['yn_1'],
        "{{yn_2}}" : meta_data['yn_2'],
    }

    table_remarks_status = {
        "{{remarks_1}}" : meta_data['remarks_1'],
        "{{remarks_2}}" : meta_data['remarks_2'],
        "{{remarks_3}}" : meta_data['remarks_3'],
        "{{remarks_4}}" : meta_data['remarks_4'],
        "{{remarks_5}}" : meta_data['remarks_5'],
        "{{remarks_6}}" : meta_data['remarks_6'],
    }

    inspected_by_status = {
        "{{inspected_by_1}}" : meta_data['inspector_name_1'],
        "{{inspected_by_2}}" : meta_data['inspector_name_2'],
    }

    main_table          = document.tables[0]

    for i in range(len(main_table.columns)):
        for cell in main_table.column_cells(i):
            for key,value in check_box_status.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key,value)
            
    for i in range(len(main_table.columns)):
        for cell in main_table.column_cells(i):
            for key,value in yes_no_status.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key,value)
            
    for i in range(len(main_table.columns)):
        for cell in main_table.column_cells(i):
            for key,value in table_remarks_status.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key,value)
            
    for i in range(len(main_table.columns)):
        for cell in main_table.column_cells(i):
            for key,value in inspected_by_status.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key,value)

    add_signature_image(
        document=document,
        meta_data=meta_data
        )

    # save document
    document.save(output_path)

    # save pdf
    """
        docx2pdf taking to much memory in executable so removed pdf for now.
    """
    #convert(output_path, output_path[:-5] + ".pdf")

    return None


def fix_spaces(meta_data):

    if "{{assembled_by}}" in meta_data.keys():
        x = 30
        l = len(meta_data["{{assembled_by}}"]) 
        if l < x:
            meta_data["{{assembled_by}}"] = meta_data["{{assembled_by}}"] +  (x - l)*" " + "\t\t"

    if "{{tested_by}}" in meta_data.keys():
        x = 30
        l = len(meta_data["{{tested_by}}"]) 
        if l < x:
            meta_data["{{tested_by}}"] = meta_data["{{tested_by}}"] +  (x - l)*" " + "\t\t"

    if "{{approved_by}}" in meta_data.keys():
        x = 30
        l = len(meta_data["{{approved_by}}"]) 
        if l < x:
            meta_data["{{approved_by}}"] = meta_data["{{approved_by}}"] +  (x - l)*" " + "\t\t"


    if "{{assembled_date}}" in meta_data.keys():
        x = 25
        l = len(meta_data["{{assembled_date}}"]) 
        if l < x:
            meta_data["{{assembled_date}}"] = meta_data["{{assembled_date}}"] +  (x - l)*" " + "\t\t"

    if "{{tested_date}}" in meta_data.keys():
        x = 25
        l = len(meta_data["{{tested_date}}"]) 
        if l < x:
            meta_data["{{tested_date}}"] = meta_data["{{tested_date}}"] +  (x - l)*" " + "\t\t"

    if "{{approved_date}}" in meta_data.keys():
        x = 25
        l = len(meta_data["{{approved_date}}"]) 
        if l < x:
            meta_data["{{approved_date}}"] = meta_data["{{approved_date}}"] +  (x - l)*" " + "\t\t"

    return meta_data


def add_signature_image(
        document                = None,
        meta_data               = {},
        assembler_signature_image_path      = "../../images/signatures/no_signature_found.png",
        tester_signature_image_path         = "../../images/signatures/no_signature_found.png",
        approver_signature_image_path       = "../../images/signatures/no_signature_found.png"
):
    
    if meta_data["assembler_name"] != "No name found":
        assembler_signature_image_path      = "../../images/signatures/{0}.png".format(meta_data["assembler_name"])

    if meta_data["tester_name"] != "No name found":
        tester_signature_image_path         = "../../images/signatures/{0}.png".format(meta_data["tester_name"])

    if meta_data["approver_name"] != "No name found":
        approver_signature_image_path       = "../../images/signatures/{0}.png".format(meta_data["approver_name"])

    # print("assembler path: ",assembler_signature_image_path)
    # print("tester path: ",tester_signature_image_path)
    # print("approver path: ",approver_signature_image_path)


    sig_height   = 0.8
    sig_width    = 2.5
    for i,paragraph in enumerate(document.paragraphs):        
        
        if i == 4:
            r = paragraph.add_run()
            r.add_picture(
                assembler_signature_image_path,
                height=Cm(sig_height),
                width=Cm(sig_width)
            )
        
        if i == 6:
            r = paragraph.add_run()
            r.add_picture(
                tester_signature_image_path,
                height=Cm(sig_height),
                width=Cm(sig_width)
            )
            
        if i == 8:
            r = paragraph.add_run()
            r.add_picture(
                approver_signature_image_path,
                height=Cm(sig_height),
                width=Cm(sig_width)
            )

    return None
